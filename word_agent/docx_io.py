"""DOCX 提取和渲染工具。

本模块尽量保留 Word 文档结构，而不是过早压平成纯文本。提取阶段会保留所有可见文本块
和弱格式信号，再由 LLM 判断哪些内容是要求、标题、正文、表格、备注或示例。
"""

from __future__ import annotations

import logging
import re
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Length

from word_agent.models import GeneratedDocument

logger = logging.getLogger(__name__)

NUMBERING_PATTERN = re.compile(
    r"^\s*((\u7b2c[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\u767e\d]+"
    r"[\u7ae0\u8282\u90e8\u5206])|([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341]+[\u3001.．])|"
    r"(\([\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+\))|"
    r"(\uff08[\u4e00\u4e8c\u4e09\u56db\u4e94\u516d\u4e03\u516b\u4e5d\u5341\d]+\uff09)|"
    r"(\d+(\.\d+)*[\u3001.．]))"
)


def _length_to_cm(value: Length | None) -> float | None:
    """将 Word 长度单位转换为便于阅读的厘米值。"""

    if value is None:
        return None
    return round(value.cm, 2)


def _trim(text: str, limit: int) -> str:
    """仅在显式配置正数限制时截断文本。"""

    if limit <= 0 or len(text) <= limit:
        return text
    return text[:limit].rstrip() + "\n...[trimmed]"


def _line_spacing_value(value: Any) -> float | str | None:
    """将 python-docx 的行距值转换为便于 JSON 序列化的数据。"""

    if value is None:
        return None
    if hasattr(value, "pt"):
        return f"{round(value.pt, 2)}pt"
    return value


def _numbering_hint(text: str) -> str | None:
    """返回可能表示标题或列表层级的开头编号标记。"""

    match = NUMBERING_PATTERN.match(text)
    if match is None:
        return None
    return match.group(0).strip()


def _paragraph_profile(paragraph: Any) -> dict[str, Any]:
    """提取段落可见文本、样式和直接格式提示。"""

    first_run = next((run for run in paragraph.runs if run.text.strip()), None)
    font = first_run.font if first_run is not None else None
    paragraph_format = paragraph.paragraph_format
    text = paragraph.text.strip()
    return {
        "text": text,
        "style": paragraph.style.name if paragraph.style is not None else "",
        "alignment": str(paragraph.alignment),
        "font_name": font.name if font is not None else None,
        "font_size_pt": font.size.pt if font is not None and font.size is not None else None,
        "bold": font.bold if font is not None else None,
        "italic": font.italic if font is not None else None,
        "line_spacing": _line_spacing_value(paragraph_format.line_spacing),
        "space_before_pt": (
            paragraph_format.space_before.pt
            if paragraph_format.space_before is not None
            else None
        ),
        "space_after_pt": (
            paragraph_format.space_after.pt
            if paragraph_format.space_after is not None
            else None
        ),
        "numbering_hint": _numbering_hint(text),
        "text_length": len(text),
    }


def _section_profile(document: DocxDocument) -> list[dict[str, Any]]:
    """汇总节级页面设置以及页眉页脚文本。"""

    sections: list[dict[str, Any]] = []
    for section in document.sections:
        sections.append(
            {
                "page_width_cm": _length_to_cm(section.page_width),
                "page_height_cm": _length_to_cm(section.page_height),
                "top_margin_cm": _length_to_cm(section.top_margin),
                "bottom_margin_cm": _length_to_cm(section.bottom_margin),
                "left_margin_cm": _length_to_cm(section.left_margin),
                "right_margin_cm": _length_to_cm(section.right_margin),
                "header_text": "\n".join(p.text for p in section.header.paragraphs if p.text.strip()),
                "footer_text": "\n".join(p.text for p in section.footer.paragraphs if p.text.strip()),
            }
        )
    return sections


def _table_profile(document: DocxDocument) -> list[dict[str, Any]]:
    """将所有表格单元格文本收集为结构化样本。"""

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables.append(
            {
                "index": table_index,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": rows,
            }
        )
    return tables


def _table_text_blocks(document: DocxDocument, start_index: int = 1) -> list[dict[str, Any]]:
    """将所有可见表格行提取为源文本块。"""

    blocks: list[dict[str, Any]] = []
    next_index = start_index
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                text = " | ".join(cells)
                blocks.append(
                    {
                        "index": next_index,
                        "source_type": "table",
                        "source": f"table_{table_index}_row_{row_index}",
                        "text": text,
                        "text_length": len(text),
                        "numbering_hint": _numbering_hint(text),
                    }
                )
                next_index += 1
    return blocks


def _header_footer_blocks(document: DocxDocument, start_index: int) -> list[dict[str, Any]]:
    """将所有可见页眉和页脚段落提取为源文本块。"""

    blocks: list[dict[str, Any]] = []
    next_index = start_index
    for section_index, section in enumerate(document.sections, start=1):
        for source_type, container in (("header", section.header), ("footer", section.footer)):
            for paragraph_index, paragraph in enumerate(container.paragraphs, start=1):
                if not paragraph.text.strip():
                    continue
                profile = _paragraph_profile(paragraph)
                profile.update(
                    {
                        "index": next_index,
                        "source_type": source_type,
                        "source": f"section_{section_index}_{source_type}_{paragraph_index}",
                    }
                )
                blocks.append(profile)
                next_index += 1
    return blocks


def _style_names(document: DocxDocument) -> list[str]:
    """返回模板中可用的段落样式名称。"""

    return sorted(
        style.name
        for style in document.styles
        if style.type == WD_STYLE_TYPE.PARAGRAPH
    )


def _paragraph_blocks(document: DocxDocument) -> list[dict[str, Any]]:
    """将所有非空正文段落提取为有序源文本块。"""

    blocks: list[dict[str, Any]] = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        if not paragraph.text.strip():
            continue
        profile = _paragraph_profile(paragraph)
        profile.update(
            {
                "index": len(blocks) + 1,
                "source_type": "paragraph",
                "source": f"paragraph_{paragraph_index}",
            }
        )
        blocks.append(profile)
    return blocks


def _ordered_text(blocks: list[dict[str, Any]]) -> str:
    """将源文本块渲染为适合提示词使用的紧凑文本列表。"""

    return "\n".join(
        f"[{item['index']:03d} {item['source_type']} style={item.get('style', '')} "
        f"num={item.get('numbering_hint')}] {item['text']}"
        for item in blocks
    )


def extract_template_profile(path: Path, max_chars: int) -> dict[str, Any]:
    """读取 DOCX 模板并返回全文、样式和版式信号。"""

    logger.info("开始提取模板 DOCX: %s", path)
    document = Document(path)
    paragraph_blocks = _paragraph_blocks(document)
    table_blocks = _table_text_blocks(document, start_index=len(paragraph_blocks) + 1)
    header_footer_blocks = _header_footer_blocks(
        document,
        start_index=len(paragraph_blocks) + len(table_blocks) + 1,
    )
    all_blocks = paragraph_blocks + table_blocks + header_footer_blocks
    logger.info(
        "模板 DOCX 提取完成，段落块=%s，表格行块=%s，页眉页脚块=%s，样式数=%s",
        len(paragraph_blocks),
        len(table_blocks),
        len(header_footer_blocks),
        len(_style_names(document)),
    )
    return {
        "path": str(path),
        "sections": _section_profile(document),
        "paragraph_styles": _style_names(document),
        "paragraph_blocks": paragraph_blocks,
        "tables": _table_profile(document),
        "all_template_text_blocks": all_blocks,
        "visible_text": _trim(_ordered_text(all_blocks), max_chars),
    }


def extract_visible_text(path: Path, max_chars: int) -> str:
    """从 DOCX 内容文件中提取所有可见段落和表格文本。"""

    profile = extract_content_profile(
        path=path,
        max_chars=max_chars,
        max_blocks=0,
        block_text_limit=0,
    )
    return profile["ordered_visible_text"]


def extract_content_profile(
    path: Path,
    max_chars: int,
    max_blocks: int,
    block_text_limit: int,
) -> dict[str, Any]:
    """提取全部内容块及弱格式证据，供 LLM 分析。"""

    logger.info("开始提取内容 DOCX: %s", path)
    document = Document(path)
    blocks = _paragraph_blocks(document)
    table_blocks = _table_text_blocks(document, start_index=len(blocks) + 1)
    blocks.extend(table_blocks)

    if block_text_limit > 0:
        for block in blocks:
            block["text"] = _trim(block["text"], block_text_limit)
            block["text_length"] = len(block["text"])

    if max_blocks > 0:
        blocks = blocks[:max_blocks]

    logger.info(
        "内容 DOCX 提取完成，内容块=%s，表格数=%s，max_blocks=%s，block_text_limit=%s",
        len(blocks),
        len(document.tables),
        max_blocks,
        block_text_limit,
    )
    return {
        "path": str(path),
        "blocks": blocks,
        "tables": _table_profile(document),
        "ordered_visible_text": _trim(_ordered_text(blocks), max_chars),
    }


def _clear_body_keep_section_properties(document: DocxDocument) -> None:
    """移除正文内容，同时保留模板中的节属性。"""

    body = document._body._element
    section_properties = body.sectPr
    for child in list(body):
        if child is not section_properties:
            body.remove(child)
    if section_properties is not None and section_properties.getparent() is None:
        body.append(section_properties)


def _has_style(document: DocxDocument, style_name: str) -> bool:
    """判断目标文档中是否存在指定段落样式。"""

    try:
        document.styles[style_name]
        return True
    except KeyError:
        return False


def _add_paragraph(document: DocxDocument, text: str, style: str, fallback_style: str) -> None:
    """在样式可用时按请求样式添加段落。"""

    selected_style = style if _has_style(document, style) else fallback_style
    if not _has_style(document, selected_style):
        selected_style = "Normal"
    document.add_paragraph(text, style=selected_style)


def write_generated_docx(
    template_path: Path,
    output_path: Path,
    generated: GeneratedDocument,
    fallback_style: str,
) -> Path:
    """基于模板将结构化生成内容渲染为新的 DOCX。"""

    logger.info("开始基于模板渲染 DOCX: template=%s output=%s", template_path, output_path)
    document = Document(template_path)
    _clear_body_keep_section_properties(document)

    if generated.title:
        title_style = "Title" if _has_style(document, "Title") else fallback_style
        _add_paragraph(document, generated.title, title_style, fallback_style)

    for paragraph in generated.paragraphs:
        if paragraph.text.strip():
            _add_paragraph(document, paragraph.text.strip(), paragraph.style, fallback_style)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    logger.info("DOCX 渲染完成，段落数=%s，输出=%s", len(generated.paragraphs), output_path)
    return output_path
