"""格式直接转换模块 — 无需 LLM，直接 Parser → GeneratedDocument → Renderer。"""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from word_agent.models import GeneratedDocument, GeneratedParagraph, ROLE_TO_DOCX_STYLE
from word_agent.parsers import UnifiedParagraph, get_parser
from word_agent.renderers import get_renderer

logger = logging.getLogger(__name__)

SUPPORTED_INPUT_EXTS = {".docx", ".md", ".markdown", ".tex", ".latex"}
SUPPORTED_OUTPUT_EXTS = {".docx", ".md", ".tex"}

_EXT_TO_TARGET = {
    "docx": ".docx",
    "md": ".md",
    "tex": ".tex",
}


def _semantic_to_role(semantic: dict | None) -> str:
    if not semantic:
        return "body"
    elem = semantic.get("element_type", "paragraph")
    if elem == "heading":
        level = semantic.get("heading_level", 1)
        if level <= 1:
            return "heading_1"
        elif level == 2:
            return "heading_2"
        else:
            return "heading_3"
    elif elem == "list_item":
        lt = semantic.get("list_type", "unordered")
        return "list_number" if lt == "ordered" else "list_bullet"
    elif elem == "code_block":
        return "code_block"
    elif elem == "blockquote":
        return "blockquote"
    return "body"


def _unified_to_generated(paragraphs: list[UnifiedParagraph]) -> GeneratedDocument:
    title = ""
    gen_paras: list[GeneratedParagraph] = []

    for p in paragraphs:
        role = _semantic_to_role(p.format.get("semantic"))
        if not title and role == "heading_1":
            title = p.text
        gen_paras.append(GeneratedParagraph(role=role, text=p.text))

    return GeneratedDocument(title=title, paragraphs=gen_paras)


def _write_docx_with_format(
    paragraphs: list[UnifiedParagraph],
    output_path: Path,
    template_path: Path | None = None,
) -> Path:
    if template_path and template_path.suffix.lower() == ".docx":
        from word_agent.docx_io import _clear_body_keep_section_properties
        doc = Document(str(template_path))
        _clear_body_keep_section_properties(doc)
    else:
        import docx as _docx
        default_tpl = Path(_docx.__file__).parent / "templates" / "default-docx-template"
        doc = Document(str(default_tpl))

    for p in paragraphs:
        role = _semantic_to_role(p.format.get("semantic"))
        style_name = ROLE_TO_DOCX_STYLE.get(role, "Normal")

        try:
            para = doc.add_paragraph(style=style_name)
        except KeyError:
            para = doc.add_paragraph(style="Normal")

        run = para.add_run(p.text)

        fmt = p.format
        if fmt.get("bold"):
            run.font.bold = True
        if fmt.get("italic"):
            run.font.italic = True
        if fmt.get("font_size_pt"):
            run.font.size = Pt(fmt["font_size_pt"])
        if fmt.get("font_name") and fmt["font_name"] != "monospace":
            run.font.name = fmt["font_name"]
        elif fmt.get("font_name") == "monospace":
            run.font.name = "Courier New"

        alignment = fmt.get("alignment")
        if alignment and alignment != "None":
            align_map = {
                "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
                "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
                "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
            }
            if alignment.upper() in align_map:
                para.alignment = align_map[alignment.upper()]

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def format_convert(input_path: Path, target_format: str) -> Path:
    target_ext = _EXT_TO_TARGET.get(target_format)
    if not target_ext:
        raise ValueError(f"不支持的目标格式: {target_format}，支持: {list(_EXT_TO_TARGET.keys())}")

    input_ext = input_path.suffix.lower()
    if input_ext not in SUPPORTED_INPUT_EXTS:
        raise ValueError(f"不支持的输入格式: {input_ext}")

    output_path = input_path.with_suffix(target_ext)

    logger.info("格式转换: %s → %s", input_path.name, output_path.name)

    parser = get_parser(input_path)
    paragraphs = parser.parse(input_path)

    if target_ext == ".docx":
        template = input_path if input_ext == ".docx" else None
        return _write_docx_with_format(paragraphs, output_path, template_path=template)
    else:
        generated = _unified_to_generated(paragraphs)
        renderer = get_renderer(output_path)
        return renderer.render(generated, output_path, template_path=None)
