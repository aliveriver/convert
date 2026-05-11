import difflib
from html import escape
from lxml import etree

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

COMPARE_KEYS = [
    "style", "alignment", "font_name", "font_size_pt",
    "bold", "italic", "line_spacing", "space_before_pt", "space_after_pt",
    "first_line_indent_pt",
]

LABEL_MAP = {
    "style": "样式",
    "alignment": "对齐",
    "font_name": "字体",
    "font_size_pt": "字号(pt)",
    "bold": "加粗",
    "italic": "斜体",
    "line_spacing": "行距",
    "space_before_pt": "段前(pt)",
    "space_after_pt": "段后(pt)",
    "first_line_indent_pt": "首行缩进(pt)",
}

_ALIGN_MAP = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "justify",
}


def _walk_style_chain(style):
    """遍历样式继承链，返回从当前样式到根的列表。"""
    visited = set()
    chain = []
    while style and style.style_id not in visited:
        visited.add(style.style_id)
        chain.append(style)
        style = style.base_style
    return chain


def _resolve_alignment(paragraph) -> str:
    """解析段落的有效对齐方式，遍历样式继承链。"""
    align = paragraph.alignment
    if align is not None:
        return _ALIGN_MAP.get(align, "left")
    for style in _walk_style_chain(paragraph.style):
        if style.paragraph_format and style.paragraph_format.alignment is not None:
            return _ALIGN_MAP.get(style.paragraph_format.alignment, "left")
    return "left"


def _resolve_font_name(paragraph) -> str | None:
    """解析字体名，优先 eastAsia，其次 ascii，遍历 run → 样式链。"""
    first_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if first_run is not None:
        rpr = first_run._element.find(qn("w:rPr"))
        if rpr is not None:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is not None:
                name = rfonts.get(qn("w:eastAsia")) or rfonts.get(qn("w:ascii")) or rfonts.get(qn("w:hAnsi"))
                if name:
                    return name
        if first_run.font.name:
            return first_run.font.name
    for style in _walk_style_chain(paragraph.style):
        if style.font and style.font.name:
            return style.font.name
        el = style.element.find(f".//{qn('w:rFonts')}")
        if el is not None:
            name = el.get(qn("w:eastAsia")) or el.get(qn("w:ascii")) or el.get(qn("w:hAnsi"))
            if name:
                return name
    return None


def _resolve_font_size(paragraph) -> float | None:
    """解析字号(pt)，遍历 run → 样式链。"""
    first_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if first_run is not None:
        if first_run.font.size is not None:
            return round(first_run.font.size.pt, 1)
        rpr = first_run._element.find(qn("w:rPr"))
        if rpr is not None:
            sz = rpr.find(qn("w:sz"))
            if sz is not None and sz.get(qn("w:val")):
                return round(int(sz.get(qn("w:val"))) / 2, 1)
            sz_cs = rpr.find(qn("w:szCs"))
            if sz_cs is not None and sz_cs.get(qn("w:val")):
                return round(int(sz_cs.get(qn("w:val"))) / 2, 1)
    for style in _walk_style_chain(paragraph.style):
        if style.font and style.font.size is not None:
            return round(style.font.size.pt, 1)
        el = style.element.find(f".//{qn('w:sz')}")
        if el is not None and el.get(qn("w:val")):
            return round(int(el.get(qn("w:val"))) / 2, 1)
    return None


def _resolve_bold(paragraph) -> bool:
    """解析加粗属性。"""
    first_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if first_run is not None:
        if first_run.font.bold is not None:
            return first_run.font.bold
    for style in _walk_style_chain(paragraph.style):
        if style.font and style.font.bold is not None:
            return style.font.bold
    return False


def _resolve_italic(paragraph) -> bool:
    """解析斜体属性。"""
    first_run = next((r for r in paragraph.runs if r.text.strip()), None)
    if first_run is not None:
        if first_run.font.italic is not None:
            return first_run.font.italic
    for style in _walk_style_chain(paragraph.style):
        if style.font and style.font.italic is not None:
            return style.font.italic
    return False


def _resolve_spacing(paragraph, attr: str):
    """解析段落间距属性，遍历样式继承链。"""
    pf = paragraph.paragraph_format
    val = getattr(pf, attr, None)
    if val is not None:
        return round(val.pt, 1)
    for style in _walk_style_chain(paragraph.style):
        if style.paragraph_format:
            val = getattr(style.paragraph_format, attr, None)
            if val is not None:
                return round(val.pt, 1)
    return None


def _resolve_line_spacing(paragraph):
    """解析行距，遍历样式继承链。"""
    pf = paragraph.paragraph_format
    if pf.line_spacing is not None:
        ls = pf.line_spacing
        if hasattr(ls, "pt"):
            return f"{round(ls.pt, 1)}pt"
        return round(ls, 2)
    for style in _walk_style_chain(paragraph.style):
        if style.paragraph_format and style.paragraph_format.line_spacing is not None:
            ls = style.paragraph_format.line_spacing
            if hasattr(ls, "pt"):
                return f"{round(ls.pt, 1)}pt"
            return round(ls, 2)
    return None


def _resolve_first_line_indent(paragraph) -> float | None:
    """解析首行缩进(pt)。"""
    pf = paragraph.paragraph_format
    if pf.first_line_indent is not None:
        return round(pf.first_line_indent.pt, 1)
    for style in _walk_style_chain(paragraph.style):
        if style.paragraph_format and style.paragraph_format.first_line_indent is not None:
            return round(style.paragraph_format.first_line_indent.pt, 1)
    return None


def _effective_format(paragraph) -> dict:
    """提取段落的有效格式（直接格式 > 样式继承链）。"""
    return {
        "style": paragraph.style.name if paragraph.style else "",
        "alignment": _resolve_alignment(paragraph),
        "font_name": _resolve_font_name(paragraph),
        "font_size_pt": _resolve_font_size(paragraph),
        "bold": _resolve_bold(paragraph),
        "italic": _resolve_italic(paragraph),
        "line_spacing": _resolve_line_spacing(paragraph),
        "space_before_pt": _resolve_spacing(paragraph, "space_before"),
        "space_after_pt": _resolve_spacing(paragraph, "space_after"),
        "first_line_indent_pt": _resolve_first_line_indent(paragraph),
    }


def _compare_format(fmt1: dict, fmt2: dict) -> list[dict]:
    """对比两个格式字典，返回结构化差异列表。"""
    changes = []
    for k in COMPARE_KEYS:
        v1 = fmt1.get(k)
        v2 = fmt2.get(k)
        if v1 != v2:
            changes.append({"key": k, "label": LABEL_MAP.get(k, k), "from": v1, "to": v2})
    return changes


def _char_diff_html(text1: str, text2: str) -> tuple[str, str]:
    """对两段文本做字符级 diff，返回带 <mark> 标签的 HTML。"""
    matcher = difflib.SequenceMatcher(None, text1, text2)
    html1_parts = []
    html2_parts = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        seg1 = escape(text1[i1:i2])
        seg2 = escape(text2[j1:j2])
        if tag == "equal":
            html1_parts.append(seg1)
            html2_parts.append(seg2)
        elif tag == "replace":
            html1_parts.append(f"<mark>{seg1}</mark>")
            html2_parts.append(f"<mark>{seg2}</mark>")
        elif tag == "delete":
            html1_parts.append(f"<mark>{seg1}</mark>")
        elif tag == "insert":
            html2_parts.append(f"<mark>{seg2}</mark>")
    return "".join(html1_parts), "".join(html2_parts)


def get_doc_diff(file1_path, file2_path):
    doc1 = Document(file1_path)
    doc2 = Document(file2_path)

    paras1 = [p for p in doc1.paragraphs if p.text.strip()]
    paras2 = [p for p in doc2.paragraphs if p.text.strip()]

    texts1 = [p.text.strip() for p in paras1]
    texts2 = [p.text.strip() for p in paras2]

    matcher = difflib.SequenceMatcher(None, texts1, texts2)
    opcodes = matcher.get_opcodes()

    diff_results = []

    for tag, i1, i2, j1, j2 in opcodes:
        if tag == "equal":
            for i, j in zip(range(i1, i2), range(j1, j2)):
                fmt1 = _effective_format(paras1[i])
                fmt2 = _effective_format(paras2[j])
                changes = _compare_format(fmt1, fmt2)

                item = {
                    "text1": texts1[i],
                    "text2": texts2[j],
                    "format1": fmt1,
                    "format2": fmt2,
                }
                if changes:
                    item["type"] = "format_diff"
                    item["format_changes"] = changes
                else:
                    item["type"] = "equal"
                diff_results.append(item)
        elif tag in ("replace", "delete", "insert"):
            text1 = "\n".join(texts1[i1:i2])
            text2 = "\n".join(texts2[j1:j2])
            html1, html2 = _char_diff_html(text1, text2)
            fmt1 = _effective_format(paras1[i1]) if i1 < i2 else {}
            fmt2 = _effective_format(paras2[j1]) if j1 < j2 else {}
            diff_results.append({
                "type": "content_diff",
                "tag": tag,
                "text1": text1,
                "text2": text2,
                "html1": html1,
                "html2": html2,
                "format1": fmt1,
                "format2": fmt2,
            })

    return diff_results
